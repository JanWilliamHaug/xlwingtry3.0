#1. remember to install "python-docx" and "xlwings" first
#2. Import libraries
import xlwings as xw
import docx
from docx import Document


#document = Document('test8.docx')
#object  "document" points to the test8 word document provided
#This function just prints out the two first lines in the word document given
document = docx.Document('test8.docx')
print(document.paragraphs[0].text)
print(document.paragraphs[1].text)

################################################################
product = "TARGEST"
coFounder1 = "Jan"
coFounder2 = "Adrian"
coFounder3 = "Stephania"
name = "NAME"
title = "TITLE"

#check if there is a keyword that you are looking for and if it is, it will replace with the name
def find_(paragraph_keyword, draft_keyword, paragraph):
    if paragraph_keyword in paragraph.text:
        print("found tag:", draft_keyword)
        #prints out "found tag:" whenever a tag is found

#going in the document.paragraphs using for loop
for paragraph in document.paragraphs:
    find_("product", product, paragraph)
    find_("Co-Founder1", coFounder1, paragraph)
    find_("Co-Founder2", coFounder2, paragraph)
    find_("Co-Founder3", coFounder3, paragraph)

    #print(paragraph.text)

#document.save(save_filename)

excelFile = xw.Book()                #Creates an empty excel file
excelFile.save('report.xlsx')        #Saves that excel file as "data1"

ws1 = excelFile.sheets['Sheet1']     #creates object called ws1 and points it to the excelFile sheet

ws1.range('A1').value = name         #Adds the string "Name" to A2
ws1.range('B1').value = title        #Adds the string "Title" to B1
ws1.range('A2').value = coFounder1   #Adds name of Co-Founder 1 to A2
ws1.range('A3').value = coFounder2   #Adds name of Co-Founder 2 to A3
ws1.range('A4').value = coFounder3   #Adds name of Co-Founder 3 to A3
ws1.range('C1').value = product      #Adds name of the product to C1

#ws1.range('A8').value = doc.paragraphs[2].text
#ws1.range('B8').value = doc.paragraphs[3].text
#ws1.range('A10:A20').value = 'Jan'
#writes my name on the cells from A2:A20
#Note: :(colon) produces one reference to all the cells between two references, including the two references.
# example'A2:A20'